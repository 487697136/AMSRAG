"""
Document parser utilities.

Extracts plain text from various file formats for RAG ingestion.
All parsers are optional — if the required library is missing,
the parser gracefully falls back to an informative error.
"""

import csv
import io
import json
from pathlib import Path
from typing import Optional


def parse_txt(content: bytes, encoding: str = "utf-8") -> str:
    return content.decode(encoding)


def parse_md(content: bytes, encoding: str = "utf-8") -> str:
    return content.decode(encoding)


def parse_json(content: bytes, encoding: str = "utf-8") -> str:
    data = json.loads(content.decode(encoding))
    if isinstance(data, list):
        parts = []
        for item in data:
            if isinstance(item, dict):
                parts.append("\n".join(f"{k}: {v}" for k, v in item.items()))
            else:
                parts.append(str(item))
        return "\n\n".join(parts)
    if isinstance(data, dict):
        return "\n".join(f"{k}: {v}" for k, v in data.items())
    return str(data)


def parse_csv(content: bytes, encoding: str = "utf-8") -> str:
    text = content.decode(encoding)
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        return ""
    return "\n".join(" | ".join(row) for row in rows)


def parse_pdf(content: bytes) -> str:
    try:
        import pdfplumber
    except ImportError:
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(content))
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text.strip())
            return "\n\n".join(pages)
        except ImportError:
            raise ImportError(
                "PDF 解析需要安装 pdfplumber 或 PyPDF2。"
                "请执行: pip install pdfplumber"
            )

    with pdfplumber.open(io.BytesIO(content)) as pdf:
        pages = []
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
    return "\n\n".join(pages)


def parse_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "DOCX 解析需要安装 python-docx。"
            "请执行: pip install python-docx"
        )

    doc = Document(io.BytesIO(content))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            paragraphs.append("\n".join(rows))

    return "\n\n".join(paragraphs)


def parse_xlsx(content: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ImportError(
            "XLSX 解析需要安装 openpyxl。"
            "请执行: pip install openpyxl"
        )

    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    sheets_text = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            sheets_text.append(f"[{sheet_name}]\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(sheets_text)


def parse_html(content: bytes, encoding: str = "utf-8") -> str:
    import re
    text = content.decode(encoding)
    text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<[^>]+>', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text


PARSER_MAP = {
    ".txt": parse_txt,
    ".md": parse_md,
    ".json": parse_json,
    ".csv": parse_csv,
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".xlsx": parse_xlsx,
    ".xls": parse_xlsx,
    ".html": parse_html,
    ".htm": parse_html,
}

TEXT_EXTENSIONS = {".txt", ".md", ".json", ".csv", ".html", ".htm"}
BINARY_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls"}


def extract_text(filename: str, content: bytes) -> str:
    """
    Extract plain text from a file's raw bytes.

    Raises ValueError if the extension is unsupported or parsing fails.
    """
    ext = Path(filename).suffix.lower()
    parser = PARSER_MAP.get(ext)
    if parser is None:
        raise ValueError(f"不支持的文件类型: {ext}")

    if ext in TEXT_EXTENSIONS:
        for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
            try:
                return parser(content, encoding=encoding)
            except (UnicodeDecodeError, TypeError):
                continue
        raise ValueError(f"无法解码文件 {filename}，请确保文件编码为 UTF-8 或 GBK。")

    return parser(content)
